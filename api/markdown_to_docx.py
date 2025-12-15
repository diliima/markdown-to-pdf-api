import markdown2
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import re
import os
import logging
from typing import Optional, List
from pathlib import Path

# Configuração de logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


class MarkdownToDocx:
    """
    Classe para converter arquivos Markdown para Word (.docx) usando python-docx.
    Compatível com ambiente serverless e Vercel.
    """
    
    def __init__(self):
        """
        Inicializa o conversor.
        """
        self.doc = None
        
    def _setup_document_styles(self):
        """
        Configura estilos customizados para o documento Word.
        """
        # Estilo para título principal (Heading 1)
        if 'Custom Title' not in [style.name for style in self.doc.styles]:
            title_style = self.doc.styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
            title_format = title_style.paragraph_format
            title_format.space_after = Pt(18)
            title_format.space_before = Pt(6)
            
            title_font = title_style.font
            title_font.name = 'Calibri'
            title_font.size = Pt(24)
            title_font.bold = True
            title_font.color.rgb = RGBColor(44, 62, 80)  # #2c3e50
        
        # Estilo para subtítulos (Heading 2)
        if 'Custom Heading 2' not in [style.name for style in self.doc.styles]:
            h2_style = self.doc.styles.add_style('Custom Heading 2', WD_STYLE_TYPE.PARAGRAPH)
            h2_format = h2_style.paragraph_format
            h2_format.space_after = Pt(12)
            h2_format.space_before = Pt(6)
            
            h2_font = h2_style.font
            h2_font.name = 'Calibri'
            h2_font.size = Pt(18)
            h2_font.bold = True
            h2_font.color.rgb = RGBColor(52, 73, 94)  # #34495e
        
        # Estilo para heading 3
        if 'Custom Heading 3' not in [style.name for style in self.doc.styles]:
            h3_style = self.doc.styles.add_style('Custom Heading 3', WD_STYLE_TYPE.PARAGRAPH)
            h3_format = h3_style.paragraph_format
            h3_format.space_after = Pt(8)
            h3_format.space_before = Pt(4)
            
            h3_font = h3_style.font
            h3_font.name = 'Calibri'
            h3_font.size = Pt(14)
            h3_font.bold = True
            h3_font.color.rgb = RGBColor(52, 73, 94)  # #34495e
        
        # Estilo para texto normal customizado
        if 'Custom Normal' not in [style.name for style in self.doc.styles]:
            normal_style = self.doc.styles.add_style('Custom Normal', WD_STYLE_TYPE.PARAGRAPH)
            normal_format = normal_style.paragraph_format
            normal_format.space_after = Pt(6)
            normal_format.line_spacing = 1.15
            
            normal_font = normal_style.font
            normal_font.name = 'Calibri'
            normal_font.size = Pt(11)
        
        # Estilo para código
        if 'Custom Code' not in [style.name for style in self.doc.styles]:
            code_style = self.doc.styles.add_style('Custom Code', WD_STYLE_TYPE.PARAGRAPH)
            code_format = code_style.paragraph_format
            code_format.space_after = Pt(6)
            code_format.left_indent = Inches(0.3)
            
            code_font = code_style.font
            code_font.name = 'Courier New'
            code_font.size = Pt(9)
        
        # Estilo para citações
        if 'Custom Quote' not in [style.name for style in self.doc.styles]:
            quote_style = self.doc.styles.add_style('Custom Quote', WD_STYLE_TYPE.PARAGRAPH)
            quote_format = quote_style.paragraph_format
            quote_format.space_after = Pt(6)
            quote_format.left_indent = Inches(0.5)
            quote_format.right_indent = Inches(0.5)
            
            quote_font = quote_style.font
            quote_font.name = 'Calibri'
            quote_font.size = Pt(11)
            quote_font.italic = True
            quote_font.color.rgb = RGBColor(93, 109, 126)  # #5d6d7e
    
    def _add_border_to_paragraph(self, paragraph):
        """
        Adiciona borda à esquerda do parágrafo (para citações).
        """
        try:
            p = paragraph._element
            pPr = p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '6')
            left.set(qn('w:space'), '1')
            left.set(qn('w:color'), '3498db')
            pBdr.append(left)
            pPr.append(pBdr)
        except Exception:
            pass  # Se falhar na formatação, continua sem borda
    
    def _parse_markdown_to_document(self, markdown_text: str):
        """
        Converte texto Markdown em elementos do documento Word.
        
        Args:
            markdown_text (str): Texto em formato Markdown
        """
        # Converte Markdown para HTML
        html = markdown2.markdown(
            markdown_text,
            extras=['fenced-code-blocks', 'tables', 'strike', 'task_list']
        )
        
        # Divide o HTML em linhas para processamento
        lines = html.split('\n')
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                i += 1
                continue
            
            # Cabeçalhos
            if line.startswith('<h1>') and line.endswith('</h1>'):
                title = self._clean_html_tags(line)
                paragraph = self.doc.add_paragraph(title, style='Custom Title')
                
            elif line.startswith('<h2>') and line.endswith('</h2>'):
                title = self._clean_html_tags(line)
                paragraph = self.doc.add_paragraph(title, style='Custom Heading 2')
                
            elif line.startswith('<h3>') and line.endswith('</h3>'):
                title = self._clean_html_tags(line)
                paragraph = self.doc.add_paragraph(title, style='Custom Heading 3')
            
            # Bloco de código
            elif line.startswith('<pre><code>'):
                # Coleta todo o bloco de código
                code_lines = []
                if line.endswith('</code></pre>'):
                    # Código em uma linha
                    code_content = line.replace('<pre><code>', '').replace('</code></pre>', '')
                    paragraph = self.doc.add_paragraph(code_content, style='Custom Code')
                else:
                    # Código em múltiplas linhas
                    code_lines.append(line.replace('<pre><code>', ''))
                    i += 1
                    while i < len(lines) and not lines[i].strip().endswith('</code></pre>'):
                        code_lines.append(lines[i])
                        i += 1
                    if i < len(lines):
                        code_lines.append(lines[i].replace('</code></pre>', ''))
                    
                    code_content = '\n'.join(code_lines)
                    paragraph = self.doc.add_paragraph(code_content, style='Custom Code')
                
            # Blockquote
            elif line.startswith('<blockquote>'):
                quote_lines = []
                if line.endswith('</blockquote>'):
                    quote_content = line.replace('<blockquote>', '').replace('</blockquote>', '')
                    paragraph = self.doc.add_paragraph(self._clean_html_tags(quote_content), style='Custom Quote')
                    self._add_border_to_paragraph(paragraph)
                else:
                    quote_lines.append(line.replace('<blockquote>', ''))
                    i += 1
                    while i < len(lines) and not lines[i].strip().endswith('</blockquote>'):
                        quote_lines.append(lines[i])
                        i += 1
                    if i < len(lines):
                        quote_lines.append(lines[i].replace('</blockquote>', ''))
                    
                    quote_content = ' '.join(quote_lines)
                    paragraph = self.doc.add_paragraph(self._clean_html_tags(quote_content), style='Custom Quote')
                    self._add_border_to_paragraph(paragraph)
            
            # Lista
            elif line.startswith('<ul>') or line.startswith('<ol>'):
                # Processa lista
                i += 1
                while i < len(lines) and not (lines[i].strip() == '</ul>' or lines[i].strip() == '</ol>'):
                    if lines[i].strip().startswith('<li>'):
                        item_text = self._clean_html_tags(lines[i].strip())
                        paragraph = self.doc.add_paragraph(item_text, style='List Bullet')
                    i += 1
            
            # Parágrafo normal
            elif line.startswith('<p>') and line.endswith('</p>'):
                text = self._clean_html_tags(line)
                paragraph = self.doc.add_paragraph(text, style='Custom Normal')
                self._apply_text_formatting(paragraph, text)
            
            # Outros elementos como texto simples
            else:
                clean_text = self._clean_html_tags(line)
                if clean_text:
                    paragraph = self.doc.add_paragraph(clean_text, style='Custom Normal')
                    self._apply_text_formatting(paragraph, line)
            
            i += 1
    
    def _apply_text_formatting(self, paragraph, original_text):
        """
        Aplica formatação de texto (negrito, itálico) ao parágrafo.
        
        Args:
            paragraph: Parágrafo do documento Word
            original_text: Texto original com tags HTML
        """
        # Limpa o parágrafo atual
        paragraph.clear()
        
        # Encontra formatações no texto original
        parts = self._split_formatted_text(original_text)
        
        for part_text, is_bold, is_italic, is_code in parts:
            run = paragraph.add_run(part_text)
            if is_bold:
                run.bold = True
            if is_italic:
                run.italic = True
            if is_code:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
    
    def _split_formatted_text(self, text):
        """
        Divide o texto em partes com suas respectivas formatações.
        
        Returns:
            List[Tuple[str, bool, bool, bool]]: (texto, negrito, itálico, código)
        """
        parts = []
        current_pos = 0
        
        # Padrões de formatação
        patterns = [
            (r'<strong>(.*?)</strong>', 'bold'),
            (r'<b>(.*?)</b>', 'bold'),
            (r'<em>(.*?)</em>', 'italic'),
            (r'<i>(.*?)</i>', 'italic'),
            (r'<code>(.*?)</code>', 'code')
        ]
        
        # Remove tags básicas primeiro
        text = self._clean_basic_tags(text)
        
        # Encontra todas as formatações
        matches = []
        for pattern, format_type in patterns:
            for match in re.finditer(pattern, text):
                matches.append((match.start(), match.end(), match.group(1), format_type))
        
        # Ordena por posição
        matches.sort()
        
        for start, end, content, format_type in matches:
            # Adiciona texto antes da formatação
            if start > current_pos:
                plain_text = text[current_pos:start]
                plain_text = re.sub(r'<[^>]+>', '', plain_text)
                if plain_text:
                    parts.append((plain_text, False, False, False))
            
            # Adiciona texto formatado
            is_bold = format_type == 'bold'
            is_italic = format_type == 'italic'
            is_code = format_type == 'code'
            parts.append((content, is_bold, is_italic, is_code))
            
            current_pos = end
        
        # Adiciona texto restante
        if current_pos < len(text):
            remaining_text = text[current_pos:]
            remaining_text = re.sub(r'<[^>]+>', '', remaining_text)
            if remaining_text:
                parts.append((remaining_text, False, False, False))
        
        return parts
    
    def _clean_basic_tags(self, text: str) -> str:
        """
        Remove apenas tags básicas, mantendo formatação.
        """
        # Remove tags básicas mas mantém formatação
        text = re.sub(r'<p.*?>', '', text)
        text = re.sub(r'</p>', '', text)
        text = re.sub(r'<li.*?>', '', text)
        text = re.sub(r'</li>', '', text)
        text = re.sub(r'<h[1-6].*?>', '', text)
        text = re.sub(r'</h[1-6]>', '', text)
        text = re.sub(r'<blockquote.*?>', '', text)
        text = re.sub(r'</blockquote>', '', text)
        return text.strip()
    
    def _clean_html_tags(self, text: str) -> str:
        """
        Remove todas as tags HTML do texto.
        
        Args:
            text (str): Texto com tags HTML
            
        Returns:
            str: Texto limpo
        """
        # Remove todas as tags HTML
        text = re.sub(r'<[^>]+>', '', text)
        return text.strip()
    
    def markdown_text_to_docx(self, markdown_text: str, output_path: str) -> bool:
        """
        Converte texto Markdown diretamente para Word (.docx).
        
        Args:
            markdown_text (str): Texto em formato Markdown
            output_path (str): Caminho onde salvar o documento gerado
            
        Returns:
            bool: True se a conversão foi bem-sucedida, False caso contrário
        """
        try:
            logger.info("Iniciando conversão de texto Markdown para Word...")
            
            # Cria um novo documento
            self.doc = Document()
            
            # Configura os estilos customizados
            self._setup_document_styles()
            
            # Converte Markdown em elementos do documento
            self._parse_markdown_to_document(markdown_text)
            
            # Salva o documento
            self.doc.save(output_path)
            
            logger.info(f"Documento Word gerado com sucesso: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Erro ao converter Markdown para Word: {str(e)}")
            return False
    
    def markdown_file_to_docx(self, markdown_file: str, output_path: Optional[str] = None) -> bool:
        """
        Converte arquivo Markdown para Word (.docx).
        
        Args:
            markdown_file (str): Caminho para o arquivo Markdown
            output_path (str, optional): Caminho onde salvar o documento
            
        Returns:
            bool: True se a conversão foi bem-sucedida, False caso contrário
        """
        try:
            # Verifica se o arquivo existe
            if not os.path.exists(markdown_file):
                logger.error(f"Arquivo Markdown não encontrado: {markdown_file}")
                return False
            
            # Define o caminho de saída se não fornecido
            if output_path is None:
                path = Path(markdown_file)
                output_path = str(path.with_suffix('.docx'))
            
            # Lê o conteúdo do arquivo Markdown
            with open(markdown_file, 'r', encoding='utf-8') as file:
                markdown_text = file.read()
            
            logger.info(f"Convertendo arquivo: {markdown_file} -> {output_path}")
            
            # Converte o texto para Word
            return self.markdown_text_to_docx(markdown_text, output_path)
            
        except Exception as e:
            logger.error(f"Erro ao processar arquivo Markdown: {str(e)}")
            return False